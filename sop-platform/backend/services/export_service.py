
import io
import logging
import re
import unicodedata
from datetime import datetime, timezone, timedelta

logger = logging.getLogger(__name__)

_STATUS_WATERMARKS = {
    "draft": "DRAFT",
    "review": "UNDER REVIEW",
    "rejected": "REJECTED",
    "archived": "ARCHIVED",
}


def _get_watermark_text(status: str) -> str | None:
   
    return _STATUS_WATERMARKS.get(status)


_UNICODE_REPLACEMENTS = {
    "\u2013": "-",   
    "\u2014": "-",    
    "\u2015": "-",    
    "\u2018": "'",    
    "\u2019": "'",    
    "\u201c": '"',    
    "\u201d": '"',    
    "\u2022": "-",   
    "\u2023": "-",    
    "\u25cf": "-",    
    "\u25a0": " ",    
    "\u25aa": " ",    
    "\u2610": " ",    
    "\u2611": "[x]",  
    "\u2612": "[ ]", 
    "\u2713": "OK",  
    "\u2714": "OK",   
    "\u2715": "X",    
    "\u2716": "X",   
    "\u26a0": "!",   
    "\u00e2": "a",    
    "\u00e2\u0080\u0099": "'",
    "\ufeff": "",     
    "\u200b": "",     
    "\u00ad": "-",    
    "\u2010": "-",
    "\u2011": "-",    
    "\u2012": "-",    
    "\u2026": "...",  
    "\u00b7": ".",  
}


def _clean(text: str) -> str:
   
    if not isinstance(text, str):
        text = str(text) if text is not None else ""

    for char, replacement in _UNICODE_REPLACEMENTS.items():
        text = text.replace(char, replacement)

    text = unicodedata.normalize("NFKD", text)

    cleaned = ""
    for ch in text:
        try:
            ch.encode("latin-1")
            cleaned += ch
        except (UnicodeEncodeError, UnicodeDecodeError):
            cleaned += " "

    cleaned = re.sub(r" {2,}", " ", cleaned)
    return cleaned





def generate_pdf(sop) -> tuple[bytes | None, str | None]:

    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm, mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, PageBreak, KeepTogether,
        )
        from reportlab.platypus.flowables import Flowable

        buffer = io.BytesIO()
        content = sop.content or {}
        doc_header = content.get("document_header", {})
        status = sop.status
        watermark_text = _get_watermark_text(status)
        ist_tz = timezone(timedelta(hours=5, minutes=30))
        generated_at = datetime.now(ist_tz).strftime("%Y-%m-%d %H:%M IST")

        styles = getSampleStyleSheet()

        style_title = ParagraphStyle(
            "SOPTitle",
            parent=styles["Title"],
            fontSize=20,
            textColor=colors.HexColor("#1E3A5F"),
            spaceAfter=6,
            fontName="Helvetica-Bold",
        )
        style_sop_number = ParagraphStyle(
            "SOPNumber",
            parent=styles["Normal"],
            fontSize=11,
            textColor=colors.HexColor("#4A6FA5"),
            spaceAfter=4,
            fontName="Helvetica-Bold",
        )
        style_section_heading = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=colors.HexColor("#1E3A5F"),
            spaceBefore=14,
            spaceAfter=6,
            fontName="Helvetica-Bold",
            borderPad=4,
        )
        style_body = ParagraphStyle(
            "SOPBody",
            parent=styles["Normal"],
            fontSize=10,
            leading=15,
            spaceAfter=4,
            fontName="Helvetica",
        )
        style_step = ParagraphStyle(
            "SOPStep",
            parent=styles["Normal"],
            fontSize=10,
            leading=15,
            spaceAfter=6,
            leftIndent=12,
            fontName="Helvetica",
        )
        style_label = ParagraphStyle(
            "SOPLabel",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#666666"),
            fontName="Helvetica-Bold",
        )
        style_center = ParagraphStyle(
            "Center",
            parent=styles["Normal"],
            alignment=TA_CENTER,
            fontSize=10,
            fontName="Helvetica",
        )

        def _on_first_page(canvas, doc):
            _draw_page_frame(canvas, doc, watermark_text, generated_at, sop, is_first=True)

        def _on_later_pages(canvas, doc):
            _draw_page_frame(canvas, doc, watermark_text, generated_at, sop, is_first=False)

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=3.5 * cm,
            bottomMargin=2.5 * cm,
            title=sop.title,
            author=doc_header.get("prepared_by", "AI SOP Platform"),
            subject="Standard Operating Procedure",
        )

        story = []

        story.append(Paragraph(_clean(sop.title), style_title))
        story.append(Paragraph(
            f"SOP Number: {_clean(sop.sop_number or 'N/A')}",
            style_sop_number,
        ))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1E3A5F")))
        story.append(Spacer(1, 6 * mm))

        approved_date = "—"
        if sop.approved_at:
            try:
                approved_date = sop.approved_at.strftime("%Y-%m-%d") if hasattr(sop.approved_at, 'strftime') else str(sop.approved_at)[:10]
            except Exception:
                approved_date = str(sop.approved_at)

        author_name = (
            sop.creator.username
            if getattr(sop, "creator", None) and getattr(sop.creator, "username", None)
            else str(sop.created_by) if hasattr(sop, "created_by") else "—"
        )

        meta_data = [
            ["Version", _clean(sop.version)],
            ["Author", _clean(author_name)],
            ["Approval Date", _clean(approved_date)],
        ]
        meta_table = Table(meta_data, colWidths=[4 * cm, 8 * cm])
        meta_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EBF0F8")),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("PADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )

        story.append(meta_table)
        story.append(Spacer(1, 8 * mm))

        if sop.description:
            story.append(Paragraph("<b>Description:</b>", style_body))
            story.append(Paragraph(_clean(sop.description), style_body))
            story.append(Spacer(1, 4 * mm))

        def add_section(title: str, body_paragraphs: list):
            section_items = [
                HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")),
                Paragraph(title, style_section_heading),
            ]
            section_items.extend(body_paragraphs)
            story.append(KeepTogether(section_items[:3]))  # keep heading with first item
            story.extend(section_items[3:])

        purpose = content.get("purpose", "")
        if purpose:
            add_section("1. Purpose", [Paragraph(_clean(str(purpose)), style_body)])

        scope = content.get("scope", "")
        if scope:
            add_section("2. Scope", [Paragraph(_clean(str(scope)), style_body)])

        definitions = content.get("definitions", [])
        if definitions:
            def_items = []
            if isinstance(definitions, list):
                for d in definitions:
                    if isinstance(d, dict):
                        def_items.append(Paragraph(
                            f"<b>{_clean(d.get('term', ''))}:</b> {_clean(d.get('definition', ''))}",
                            style_body,
                        ))
                    else:
                        def_items.append(Paragraph(_clean(str(d)), style_body))
            else:
                def_items.append(Paragraph(_clean(str(definitions)), style_body))
            add_section("3. Definitions", def_items)

        responsibilities = content.get("responsibilities", [])
        if responsibilities:
            resp_items = []
            if isinstance(responsibilities, list):
                for r in responsibilities:
                    if isinstance(r, dict):
                        role_name = _clean(r.get("role", "Role"))
                        resp_data = r.get("responsibilities") or r.get("responsibility")
                        
                        if isinstance(resp_data, list):
                            resp_items.append(Paragraph(f"<b>{role_name}:</b>", style_body))
                            for item in resp_data:
                                resp_items.append(Paragraph(f"• {_clean(str(item))}", style_step))
                        else:
                            resp_items.append(Paragraph(f"<b>{role_name}:</b> {_clean(str(resp_data or ''))}", style_body))
                    else:
                        resp_items.append(Paragraph(_clean(str(r)), style_body))
            else:
                resp_items.append(Paragraph(_clean(str(responsibilities)), style_body))
            add_section("4. Responsibilities", resp_items)

        procedure = content.get("procedure", [])
        if procedure:
            proc_items = []
            if isinstance(procedure, list):
                for i, step in enumerate(procedure, 1):
                    if isinstance(step, dict):
                        step_title = _clean(step.get("step_title") or step.get("title") or f"Step {i}")
                        step_action = _clean(step.get("action") or step.get("description") or "")
                        step_owner = _clean(step.get("performed_by", ""))
                        step_note = _clean(step.get("data_integrity_note") or step.get("note") or "")
                        step_verification = _clean(step.get("verification", ""))
                        step_exception = _clean(step.get("exception_handling", ""))
                        
                        title_text = f"<b>Step {step.get('step_number', step.get('step', i))}: {step_title}</b>"
                        if step_owner:
                            title_text += f" ({step_owner})"
                        proc_items.append(Paragraph(title_text, style_step))
                        
                        if step_action:
                            proc_items.append(Paragraph(step_action, style_step))
                        if step_verification:
                            proc_items.append(Paragraph(f"<b>Verification:</b> {step_verification}", style_step))
                        if step_note:
                            proc_items.append(Paragraph(f"<i>Data Integrity Note: {step_note}</i>", style_step))
                        if step_exception:
                            proc_items.append(Paragraph(f"<i>Exception: {step_exception}</i>", style_step))
                    else:
                        proc_items.append(Paragraph(f"{i}. {_clean(str(step))}", style_step))
            else:
                proc_items.append(Paragraph(_clean(str(procedure)), style_body))
            add_section("5. Procedure", proc_items)

        references = content.get("references", [])
        if references:
            ref_items = []
            if isinstance(references, list):
                for ref in references:
                    if isinstance(ref, dict):
                        ref_str = _clean(ref.get("title", str(ref)))
                        if ref.get("document_id"):
                            ref_str += f" [{_clean(ref['document_id'])}]"
                        ref_items.append(Paragraph(f"- {ref_str}", style_body))
                    else:
                        ref_items.append(Paragraph(f"- {_clean(str(ref))}", style_body))
            else:
                ref_items.append(Paragraph(_clean(str(references)), style_body))
            add_section("6. References", ref_items)

        revision_history = content.get("revision_history", [])
        if revision_history and isinstance(revision_history, list):
            rev_table_data = [["Version", "Date", "Author", "Summary"]]
            for rev in revision_history:
                if isinstance(rev, dict):
                    rev_table_data.append([
                        _clean(str(rev.get("version", ""))),
                        _clean(str(rev.get("date", ""))),
                        _clean(str(rev.get("author", ""))),
                        _clean(str(rev.get("summary", ""))),
                    ])
            if len(rev_table_data) > 1:
                rev_table = Table(rev_table_data, colWidths=[2 * cm, 3 * cm, 4 * cm, 11 * cm])
                rev_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F8FF")]),
                    ("PADDING", (0, 0), (-1, -1), 5),
                ]))
                add_section("7. Revision History", [rev_table])

        approval_block = content.get("approval_block", {})
        if approval_block:
            appr_data = [
                ["Role", "Name", "Signature", "Date"],
                ["Prepared By", _clean(approval_block.get("prepared_by", "")), "_________________", _clean(approval_block.get("approval_date", ""))],
                ["Reviewed By", _clean(approval_block.get("reviewed_by", "")), "_________________", _clean(approval_block.get("next_review_date", ""))],
                ["Approved By", _clean(approval_block.get("approved_by", "")), "_________________", _clean(approval_block.get("approval_date", ""))],
            ]
            appr_table = Table(appr_data, colWidths=[4 * cm, 5 * cm, 5 * cm, 6 * cm])
            appr_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A5F")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("PADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            add_section("8. Approval Block", [appr_table])



        doc.build(story, onFirstPage=_on_first_page, onLaterPages=_on_later_pages)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        logger.info("PDF generated for sop_id=%s (%d bytes)", sop.id, len(pdf_bytes))
        return pdf_bytes, None

    except Exception:
        logger.exception("PDF generation failed for sop_id=%s", sop.id if sop else "?")
        return None, "PDF generation failed. Please try again."


def _draw_page_frame(canvas, doc, watermark_text, generated_at, sop, is_first: bool):
    
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm

    canvas.saveState()
    width, height = doc.pagesize

    canvas.setFillColor(colors.HexColor("#1E3A5F"))
    canvas.rect(0, height - 1.8 * cm, width, 1.8 * cm, fill=True, stroke=False)
    canvas.setFillColor(colors.white)
    canvas.setFont("Helvetica-Bold", 10)
    canvas.drawString(2 * cm, height - 1.2 * cm, "AI SOP COMPLIANCE PLATFORM")
    canvas.setFont("Helvetica", 9)
    canvas.drawRightString(width - 2 * cm, height - 1.2 * cm, f"{sop.sop_number or 'SOP'} (v{sop.version})")

    canvas.setFillColor(colors.HexColor("#F0F4FA"))
    canvas.rect(0, 0, width, 1.5 * cm, fill=True, stroke=False)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.setFont("Helvetica", 8)
    canvas.drawString(2 * cm, 0.6 * cm, f"Generated: {generated_at}  |  CONFIDENTIAL - FOR INTERNAL USE ONLY")
    canvas.drawRightString(width - 2 * cm, 0.6 * cm, f"Page {doc.page}")

    if watermark_text:
        canvas.setFont("Helvetica-Bold", 60)
        canvas.setFillColor(colors.HexColor("#E8E8E8"))
        canvas.setFillAlpha(0.35)
        canvas.saveState()
        canvas.translate(width / 2, height / 2)
        canvas.rotate(45)
        canvas.drawCentredString(0, 0, watermark_text)
        canvas.restoreState()

    canvas.restoreState()



def generate_docx(sop) -> tuple[bytes | None, str | None]:
   
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        content = sop.content or {}
        doc_header_data = content.get("document_header", {})
        status = sop.status
        watermark_text = _get_watermark_text(status)

        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        def add_heading(text: str, level: int = 1):
            h = doc.add_heading(text, level=level)
            if level == 1:
                h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
                h.runs[0].font.size = Pt(16)
            elif level == 2:
                h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
                h.runs[0].font.size = Pt(12)
            return h

        def add_body_para(text: str, bold: bool = False) -> None:
            p = doc.add_paragraph(text)
            p.runs[0].bold = bold if p.runs else False
            p.paragraph_format.space_after = Pt(6)

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_run = title_para.add_run(_clean(sop.title))
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

        sub = doc.add_paragraph(
            f"SOP Number: {_clean(sop.sop_number or 'N/A')}  |  Status: {sop.status.upper()}  |  Version: {sop.version}"
        )
        sub.runs[0].font.size = Pt(10)
        sub.runs[0].font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

        doc.add_paragraph()  # spacer

        if watermark_text:
            notice = doc.add_paragraph(f"[!] THIS DOCUMENT IS {watermark_text} - NOT FOR OFFICIAL USE")
            notice.runs[0].bold = True
            notice.runs[0].font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
            notice.alignment = WD_ALIGN_PARAGRAPH.CENTER

        author_name = (
            sop.creator.username
            if getattr(sop, "creator", None) and getattr(sop.creator, "username", None)
            else str(sop.created_by) if hasattr(sop, "created_by") else "-"
        )
        approved_date = str(sop.approved_at)[:10] if sop.approved_at else "-"
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.style = "Table Grid"
        meta_rows = [
            ("Version", _clean(sop.version)),
            ("Author", _clean(author_name)),
            ("Approval Date", _clean(approved_date)),
        ]
        for i, (label, value) in enumerate(meta_rows):
            row = meta_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        if sop.description:
            doc.add_paragraph(f"Description: {_clean(sop.description)}")
            doc.add_paragraph()

        def section_heading(number: int, title: str):
            doc.add_paragraph()
            add_heading(f"{number}. {title}", level=2)

        purpose = content.get("purpose", "")
        if purpose:
            section_heading(1, "Purpose")
            doc.add_paragraph(_clean(str(purpose)))

        scope = content.get("scope", "")
        if scope:
            section_heading(2, "Scope")
            doc.add_paragraph(_clean(str(scope)))

        definitions = content.get("definitions", [])
        if definitions:
            section_heading(3, "Definitions")
            if isinstance(definitions, list):
                for d in definitions:
                    if isinstance(d, dict):
                        p = doc.add_paragraph()
                        p.add_run(f"{_clean(d.get('term', ''))}: ").bold = True
                        p.add_run(_clean(d.get("definition", "")))
                    else:
                        doc.add_paragraph(_clean(str(d)))

        responsibilities = content.get("responsibilities", [])
        if responsibilities:
            section_heading(4, "Responsibilities")
            if isinstance(responsibilities, list):
                for r in responsibilities:
                    if isinstance(r, dict):
                        p = doc.add_paragraph()
                        run = p.add_run(f"{_clean(r.get('role', ''))}: ")
                        run.bold = True
                        p.add_run(_clean(r.get("responsibility", "")))
                    else:
                        doc.add_paragraph(_clean(str(r)))

        procedure = content.get("procedure", [])
        if procedure:
            section_heading(5, "Procedure")
            if isinstance(procedure, list):
                for i, step in enumerate(procedure, 1):
                    if isinstance(step, dict):
                        p = doc.add_paragraph()
                        run = p.add_run(f"Step {step.get('step', i)}: {_clean(step.get('title', ''))}")
                        run.bold = True
                        if step.get("description"):
                            doc.add_paragraph(_clean(step["description"]))
                        if step.get("note"):
                            note_p = doc.add_paragraph(f"Note: {_clean(step['note'])}")
                            if note_p.runs:
                                note_p.runs[0].italic = True
                    else:
                        doc.add_paragraph(f"{i}. {_clean(str(step))}")

        references = content.get("references", [])
        if references:
            section_heading(6, "References")
            if isinstance(references, list):
                for ref in references:
                    if isinstance(ref, dict):
                        ref_str = _clean(ref.get("title", str(ref)))
                        if ref.get("document_id"):
                            ref_str += f" [{_clean(ref['document_id'])}]"
                        doc.add_paragraph(f"- {ref_str}")
                    else:
                        doc.add_paragraph(f"- {_clean(str(ref))}")

        revision_history = content.get("revision_history", [])
        if revision_history and isinstance(revision_history, list) and len(revision_history) > 0:
            section_heading(7, "Revision History")
            rev_table = doc.add_table(rows=1, cols=4)
            rev_table.style = "Table Grid"
            hdr_cells = rev_table.rows[0].cells
            for i, h in enumerate(["Version", "Date", "Author", "Summary"]):
                hdr_cells[i].text = h
                hdr_cells[i].paragraphs[0].runs[0].bold = True
            for rev in revision_history:
                if isinstance(rev, dict):
                    row_cells = rev_table.add_row().cells
                    row_cells[0].text = _clean(str(rev.get("version", "")))
                    row_cells[1].text = _clean(str(rev.get("date", "")))
                    row_cells[2].text = _clean(str(rev.get("author", "")))
                    row_cells[3].text = _clean(str(rev.get("summary", "")))

        approval_block = content.get("approval_block", {})
        if approval_block:
            section_heading(8, "Approval Block")
            appr_table = doc.add_table(rows=4, cols=4)
            appr_table.style = "Table Grid"
            appr_header = appr_table.rows[0].cells
            for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
                appr_header[i].text = h
                appr_header[i].paragraphs[0].runs[0].bold = True
            appr_rows_data = [
                ("Prepared By", _clean(approval_block.get("prepared_by", "")), "_________________", _clean(approval_block.get("approval_date", ""))),
                ("Reviewed By", _clean(approval_block.get("reviewed_by", "")), "_________________", _clean(approval_block.get("next_review_date", ""))),
                ("Approved By", _clean(approval_block.get("approved_by", "")), "_________________", _clean(approval_block.get("approval_date", ""))),
            ]
            for i, (role, name, sig, date) in enumerate(appr_rows_data, 1):
                cells = appr_table.rows[i].cells
                cells[0].text = role
                cells[1].text = name
                cells[2].text = sig
                cells[3].text = date



        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        logger.info("DOCX generated for sop_id=%s (%d bytes)", sop.id, len(docx_bytes))
        return docx_bytes, None

    except Exception:
        logger.exception("DOCX generation failed for sop_id=%s", sop.id if sop else "?")
        return None, "DOCX generation failed. Please try again."
